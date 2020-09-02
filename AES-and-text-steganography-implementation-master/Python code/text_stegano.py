import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
# Styles to which words will be changed according to the bits pair
def style_set(output):
    obj_styles = output.styles

    obj_charstyle = obj_styles.add_style('Style_00', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(9)
    obj_font.name = 'Arial'

    obj_charstyle = obj_styles.add_style('Style_01', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(10)
    obj_font.name = 'Arial'

    obj_charstyle = obj_styles.add_style('Style_10', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(10)
    obj_font.name = 'Cambria'

    obj_charstyle = obj_styles.add_style('Style_11', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Arial'

    obj_charstyle = obj_styles.add_style('Style_0', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(9)
    obj_font.name = 'Calibri'

    obj_charstyle = obj_styles.add_style('Style_1', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(10)
    obj_font.name = 'Calibri'

    obj_charstyle = obj_styles.add_style('Default', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri'

# Takes input as string and returns binary
def str2bin(text):
    binary_val = bin(int.from_bytes(text.encode(), 'big'))
    return binary_val

# Takes input as binary and returns string
def bin2str(binary_val):
    binary_val = "0b" + binary_val
    n = int(binary_val, 2)
    text = n.to_bytes((n.bit_length() + 7) // 8, 'big').decode()
    return text

# Adding styled words (according to bits pair) into the output file
def encryptbinaryintotext(word, para, data):
    if data == "0":
        para.add_run(word, style='Style_0')

    if data == "1":
        para.add_run(word, style='Style_1')

    if data == "00":
        para.add_run(word, style='Style_00')

    if data == "01":
        para.add_run(word, style='Style_01')

    if data == "10":
        para.add_run(word, style='Style_10')

    if data == "11":
        para.add_run(word, style='Style_11')

# Takes input string and creates encrypted docx file
def encrypt(text):
    while True:
        try:
            dir_cover = input("Enter the directory of docx file to be used as cover: ")
            inputfile = docx.Document(dir_cover)
            error = 0
        except:
            print("Could not find file in the input directory ... try again")
            error = 1

        if error == 0:
            break
    output_name = input("Enter encrypted file's name you want as output (without .docx): ")

    if output_name[-5:] != ".docx" :
        output_name = output_name + ".docx"

    output = docx.Document()
    binary_data = str2bin(text)
    style_set(output)
    x = output.add_paragraph()
    j = 2
    word = ""
    try:
        for p in inputfile.paragraphs:
            for r in p.runs:
                for i in r.text:
                    word = word + i
                    if i == ' ' or i == '.':
                        data = binary_data[j:j+2]
                        j = j + 2
                        if data:
                            encryptbinaryintotext(word, x, data)
                            word = ""
                        else:
                            x.add_run(word, style='Default')
                            word = ""
    except:
        print("Cover docx file selected is too small to encrypt all data, select another file")

    output.save(output_name)

# This function returns base64 string after decryption
def decrypt():
    while True:
        try:
            dir_cover = input("Enter the directory of file to be decrypted : ")
            inputfile = docx.Document(dir_cover)
            error = 0
        except:
            print("Could not find docx file in the input directory ... try again")
            error = 1
        if error == 0:
            break
    data = ""
    for p in inputfile.paragraphs:
        for r in p.runs:
            if r.style.font.name == 'Cambria':
                data = data + "10"
            elif r.style.font.name == 'Calibri':
                x = r.style.font.size.pt
                if x == 9:
                    data = data + "0"
                if x == 10:
                    data = data + "1"
                if x == 11:
                    return bin2str(data)
            elif r.style.font.name == 'Arial':
                x = r.style.font.size.pt
                if x == 9:
                    data = data + "00"
                if x == 10:
                    data = data + "01"
                if x == 11:
                    data = data + "11"
    return data

# To apply only Text-Steganography
def main():
    while True:
        action = input("Do you want to encrypt or decrypt the file (d/e) or exit: ")
        if action == 'e':
            text = input("Enter the data to encrypt: ")
            encrypt(text)
        elif action == 'd':
            text = decrypt()
            print("Decrypted code is: " + text)
        elif action == 'exit':
            break
        else:
            print("Invalid input try again")



if __name__ == '__main__':
    main()
